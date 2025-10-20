# rag_utils.py
import os
import pdfplumber
import docx
import requests
from bs4 import BeautifulSoup
from readability import Document
from sentence_transformers import SentenceTransformer
# Compatibility fix for LangChain text splitter
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ModuleNotFoundError:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
import faiss
import numpy as np
import pickle

# Embedding model setup
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"
EMBEDDINGS_DIM = 384  # matches the model above

class Ingestor:
    def __init__(self, model_name: str = EMBEDDING_MODEL_NAME, index_path: str = "data/index.faiss"):
        self.model = SentenceTransformer(model_name)
        self.index_path = index_path
        self.index = None
        self.metadata = []

    # -------- TEXT EXTRACTION --------
    def extract_text_from_pdf(self, path: str) -> str:
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                texts.append(page.extract_text() or "")
        return "\n".join(texts)

    def extract_text_from_docx(self, path: str) -> str:
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])

    def fetch_url_text(self, url: str) -> str:
        r = requests.get(url, timeout=15)
        r.raise_for_status()
        doc = Document(r.text)
        html = doc.summary()
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n")

    # -------- CHUNKING --------
    def chunk_text(self, text: str, chunk_size: int = 800, chunk_overlap: int = 200):
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", " "]
        )
        return splitter.split_text(text)

    # -------- EMBEDDINGS --------
    def embed_texts(self, texts):
        return self.model.encode(texts, show_progress_bar=True, convert_to_numpy=True)

    # -------- INDEXING & RETRIEVAL --------
    def create_or_load_index(self):
        if os.path.exists(self.index_path):
            self.index = faiss.read_index(self.index_path)
            meta_path = self.index_path + ".meta"
            if os.path.exists(meta_path):
                with open(meta_path, "rb") as f:
                    self.metadata = pickle.load(f)
        else:
            self.index = faiss.IndexFlatL2(EMBEDDINGS_DIM)
            self.metadata = []

    def add_texts(self, texts, metadatas):
        embs = self.embed_texts(texts)
        if self.index is None:
            self.create_or_load_index()
        self.index.add(embs)
        self.metadata.extend(metadatas)
        faiss.write_index(self.index, self.index_path)
        with open(self.index_path + ".meta", "wb") as f:
            pickle.dump(self.metadata, f)

    def query(self, query_text, k: int = 5):
        q_emb = self.model.encode([query_text], convert_to_numpy=True)
        D, I = self.index.search(q_emb, k)
        results = []
        for idx in I[0]:
            if idx < len(self.metadata):
                results.append(self.metadata[idx])
        return results
